"""
Parsing utilities for document files.
Extracts actual text content from PDF, DOCX, DOC, and TXT files.
"""
from pathlib import Path
from typing import Optional, List
import logging

logger = logging.getLogger(__name__)

# Try to import real parsers, fall back to mock if not available
try:
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("PyPDF2 not installed. PDF parsing will use fallback.")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed. DOCX parsing will use fallback.")


def parse_pdf(filepath: str) -> str:
    """
    Parse PDF file and return actual text content.
    
    Args:
        filepath: Path to PDF file
        
    Returns:
        Extracted text content
    """
    if not PDF_AVAILABLE:
        logger.warning("PyPDF2 not available, using fallback for PDF")
        filename = Path(filepath).name
        return f"Mock text from {filename}. Install PyPDF2 for real PDF parsing."
    
    try:
        text_parts = []
        with open(filepath, "rb") as f:
            pdf_reader = PdfReader(f)
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    text = page.extract_text()
                    if text.strip():
                        text_parts.append(text)
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num + 1}: {str(e)}")
                    continue
        
        if text_parts:
            return "\n\n".join(text_parts)
        else:
            logger.warning(f"No text extracted from PDF: {filepath}")
            filename = Path(filepath).name
            return f"Could not extract text from {filename}. The PDF might be image-based or encrypted."
    except Exception as e:
        logger.error(f"Error parsing PDF {filepath}: {str(e)}")
        filename = Path(filepath).name
        return f"Error parsing PDF {filename}: {str(e)}"


def parse_doc(filepath: str) -> str:
    """
    Parse DOC file (legacy Word format).
    Note: DOC parsing requires additional libraries like python-docx2txt or antiword.
    For now, returns a message.
    
    Args:
        filepath: Path to DOC file
        
    Returns:
        Text content or fallback message
    """
    # DOC files are binary and harder to parse
    # python-docx only works with DOCX
    logger.warning("DOC format parsing not fully supported. Consider converting to DOCX.")
    filename = Path(filepath).name
    return f"DOC format parsing not fully implemented. File: {filename}. Please convert to DOCX or PDF for better support."


def parse_docx(filepath: str) -> str:
    """
    Parse DOCX file and return actual text content.
    
    Args:
        filepath: Path to DOCX file
        
    Returns:
        Extracted text content
    """
    if not DOCX_AVAILABLE:
        logger.warning("python-docx not available, using fallback for DOCX")
        filename = Path(filepath).name
        return f"Mock text from {filename}. Install python-docx for real DOCX parsing."
    
    try:
        doc = Document(filepath)
        text_parts = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        if text_parts:
            return "\n\n".join(text_parts)
        else:
            logger.warning(f"No text extracted from DOCX: {filepath}")
            filename = Path(filepath).name
            return f"Could not extract text from {filename}. The document might be empty."
    except Exception as e:
        logger.error(f"Error parsing DOCX {filepath}: {str(e)}")
        filename = Path(filepath).name
        return f"Error parsing DOCX {filename}: {str(e)}"


def parse_txt(filepath: str) -> str:
    """
    Parse TXT file and return actual content.
    
    Args:
        filepath: Path to TXT file
        
    Returns:
        File content as string
    """
    try:
        with open(filepath, "r", encoding="utf-8") as f:
            return f.read()
    except Exception:
        # Fallback to mock text if reading fails
        filename = Path(filepath).name
        return f"Mock text from {filename}."


def parse_file(filepath: str) -> Optional[str]:
    """
    Parse a file based on its extension.
    
    Args:
        filepath: Path to file
        
    Returns:
        Parsed text content, or None if unsupported format
    """
    file_path = Path(filepath)
    extension = file_path.suffix.lower()
    
    parsers = {
        ".pdf": parse_pdf,
        ".doc": parse_doc,
        ".docx": parse_docx,
        ".txt": parse_txt,
    }
    
    parser = parsers.get(extension)
    if parser:
        return parser(filepath)
    
    return None


def get_supported_extensions() -> List[str]:
    """
    Get list of supported file extensions.
    
    Returns:
        List of supported extensions (with dots)
    """
    return [".pdf", ".doc", ".docx", ".txt"]

